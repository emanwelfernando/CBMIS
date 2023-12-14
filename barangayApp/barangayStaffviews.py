from django.shortcuts import render, redirect, get_object_or_404
from django.contrib import messages
from django.db.models import Q, Sum, Count, Avg, Case, When, Value, F
from django.db.models.functions import Round
from django.http import HttpResponseRedirect, JsonResponse, HttpResponseServerError, HttpResponse, FileResponse
from django.views import View
from django.core.files.storage import default_storage
from django.core.paginator import Paginator
from django.urls import reverse
from django.contrib.auth.decorators import login_required, permission_required
from django.contrib.auth.mixins import LoginRequiredMixin, PermissionRequiredMixin
from django.contrib.auth.models import Permission
from django.contrib.gis.db import models
from django.utils import timezone
from django.conf import settings

from .forms import *
from .models import *
from .filters import *

import json
import os
from io import BytesIO

from django.template.loader import get_template
from xhtml2pdf import pisa
from pptx import Presentation
from docx import Document
from datetime import date
from datetime import datetime

from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.lib import colors

import openpyxl
from openpyxl.utils.dataframe import dataframe_to_rows
from openpyxl.utils import get_column_letter
from django.http import HttpResponse
import pandas as pd
from openpyxl import Workbook

@login_required
@permission_required('barangayApp.can_view_barangay_staff_views', raise_exception=True)
def dashboard(request):
    user = request.user
    barangay = user.barangay
    
    if not barangay:
        return HttpResponse("You are not associated with any barangay.")
    
    total_residents = Resident.objects.filter(barangay=barangay).count()
    total_households = Household.objects.filter(barangay=barangay).count()
    total_members = Household.objects.filter(barangay=barangay).aggregate(total_members=Count('members'))['total_members']
    average_members = Household.objects.filter(barangay=barangay).aggregate(average_members=Round(Avg('num_members'), 2))['average_members']

    household_data = Household.objects.filter(barangay=barangay).aggregate(
        num_senior_citizens=Count('num_senior_citizens', distinct=True),
        num_pregnant_lactating_mothers=Count('num_pregnant_lactating_mothers', distinct=True),
        num_beneficiaries_with_disability=Count('num_beneficiaries_with_disability', distinct=True),
        num_registered_voters=Count('num_registered_voters', distinct=True),
        num_members=Count('members', distinct=True)
    )

    total_senior_citizens = Household.objects.filter(barangay=barangay).aggregate(total_senior_citizens=Count('num_senior_citizens'))['total_senior_citizens']
    total_pregnant_lactating_mothers = Household.objects.filter(barangay=barangay).aggregate(total_pregnant_lactating_mothers=Count('num_pregnant_lactating_mothers'))['total_pregnant_lactating_mothers']
    total_beneficiaries_with_disability = Household.objects.filter(barangay=barangay).aggregate(total_beneficiaries_with_disability=Count('num_beneficiaries_with_disability'))['total_beneficiaries_with_disability']
    total_registered_voters = Household.objects.filter(barangay=barangay).aggregate(total_registered_voters=Count('num_registered_voters'))['total_registered_voters']

    gender_distribution = Resident.objects.filter(barangay=barangay).values('sex').annotate(count=Count('sex'))
    gender_labels = [entry['sex'] for entry in gender_distribution]
    gender_counts = [entry['count'] for entry in gender_distribution]

    age_distribution = Resident.objects.filter(barangay=barangay) \
        .annotate(age_group=Case(
            When(age__lt=18, then=Value('0-17')),
            When(age__range=[18, 25], then=Value('18-25')),
            When(age__range=[26, 35], then=Value('26-35')),
            When(age__range=[36, 50], then=Value('36-50')),
            When(age__range=[51, 60], then=Value('51-60')),
            When(age__gte=61, then=Value('60+')),
            default=Value('Unknown'),
            output_field=models.CharField(),
        )) \
        .values('age_group') \
        .annotate(count=Count('id'))

    age_labels = [item['age_group'] for item in age_distribution]
    age_counts = [item['count'] for item in age_distribution]

    worker_counts = Economic.objects.filter(resident__barangay=barangay) \
        .values('class_of_worker') \
        .annotate(count=Count('class_of_worker')) \
        .order_by('-count')[:2] 

    chart_data = {
        'labels': [item['class_of_worker'] for item in worker_counts],
        'data': [item['count'] for item in worker_counts],
    }

    chart_data_json = json.dumps(chart_data)


    context = {
        "total_residents": total_residents, 
        "total_households": total_households, 
        "total_members": total_members,
        "total_senior_citizens": total_senior_citizens,
        "total_pregnant_lactating_mothers": total_pregnant_lactating_mothers,
        "total_beneficiaries_with_disability": total_beneficiaries_with_disability,
        "total_registered_voters": total_registered_voters,

        "gender_labels": gender_labels,
        "gender_counts": gender_counts,
        "age_labels": age_labels,
        "age_counts": age_counts,
        "average_members": average_members,
        "chart_data": chart_data_json,
        }
    return render(request, 'barangayStaff/dashboard.html', context)

@login_required
@permission_required('barangayApp.can_view_barangay_staff_views', raise_exception=True)
def household(request, page=1):
    user = request.user
    barangay = user.barangay
    if not barangay:
        return HttpResponse("You are not associated with any barangay.")

    households = Household.objects.filter(barangay=barangay).order_by('household_head__last_name')
    household_filter = HouseholdFilters(request.GET, queryset=households)
    households = household_filter.qs

    household_data = []

    for household in households:
        first_name = household.household_head.first_name if household.household_head else 'No Head Assigned'
        middle_initial = household.household_head.middle_name[0] if household.household_head.middle_name else ''
        last_name = household.household_head.last_name if household.household_head else ''
        suffix = household.household_head.suffix if household.household_head else ''

        full_name = f"{last_name}, {first_name} {middle_initial}. {suffix}"

        household_data.append({
            'latitude': household.latitude,
            'longitude': household.longitude,
            'id': household.id,
            'head_name': full_name
        })

    household_data_json = json.dumps(household_data)

    p = Paginator(households, 10)  
    page_num = request.GET.get('page', 1)
    page = p.page(page_num)
    start_index = (page.number - 1) * p.per_page

    barangay_geojson = barangay.geom.geojson if barangay and barangay.geom else None

    context = {
        'households': page,
        'household_filter': household_filter,
        'start_index': start_index,
        'household_data_json': household_data_json,
        'barangay_geojson': barangay_geojson
    }
    return render(request, 'barangayStaff/household.html', context)

@login_required
@permission_required('barangayApp.can_view_barangay_staff_views', raise_exception=True)
def delete_household(request, pk):
    household = Household.objects.get(id=pk)
    household.delete()
    return redirect('barangay-staff-household')


@login_required
@permission_required('barangayApp.can_view_barangay_staff_views', raise_exception=True)
def resident(request):
    user = request.user
    barangay = user.barangay
    if not barangay:
        return HttpResponse("You are not associated with any barangay.")

    residents = Resident.objects.filter(barangay=barangay).order_by('last_name')
    
    # Apply filters if they exist in the request
    if request.GET:
        resident_filter = ResidentFilters(request.GET, queryset=residents)
        residents = resident_filter.qs
    else:
        resident_filter = ResidentFilters(queryset=residents)

    p = Paginator(residents, 10)  
    page_num = request.GET.get('page', 1)
    page = p.page(page_num)
    start_index = (page.number - 1) * p.per_page

    context = {
        'residents': page, 
        'resident_filter': resident_filter, 
        'start_index': start_index,
    }
    return render(request, 'barangayStaff/resident.html', context)

@login_required
@permission_required('barangayApp.can_view_barangay_staff_views', raise_exception=True)
def document(request):
    user = request.user
    barangay = user.barangay
    if not barangay:
        return HttpResponse("You are not associated with any barangay.")

    residents = Resident.objects.filter(barangay=barangay).order_by('last_name')
    resident_filter = ResidentFilters(request.GET, queryset=residents)
    residents = resident_filter.qs

    p = Paginator(residents, 10)  
    page_num = request.GET.get('page', 1)
    page = p.page(page_num)
    start_index = (page.number - 1) * p.per_page

    # Fetch the list of uploaded Word documents for clearance for the current barangay
    clearance_folder_path = os.path.join(settings.MEDIA_ROOT, 'doc_templates', str(barangay.id), 'clearance')
    clearance_files = []
    if os.path.exists(clearance_folder_path):
        clearance_files = [f for f in os.listdir(clearance_folder_path) if f.endswith('.docx')]

    indigency_folder_path = os.path.join(settings.MEDIA_ROOT, 'doc_templates', str(barangay.id), 'indigency')
    indigency_files = []
    if os.path.exists(indigency_folder_path):
        indigency_files = [f for f in os.listdir(indigency_folder_path) if f.endswith('.docx')]

    if request.method == 'POST':
        form = UploadClearanceForm(request.POST, request.FILES)
        if form.is_valid():
            uploaded_file = request.FILES['clearance_file']
            
            # Save the file to the user's barangay folder in the media directory
            folder_path = os.path.join(settings.MEDIA_ROOT, 'doc_templates', str(barangay.id), 'clearance')
            os.makedirs(folder_path, exist_ok=True)
            
            with open(os.path.join(folder_path, uploaded_file.name), 'wb') as destination:
                for chunk in uploaded_file.chunks():
                    destination.write(chunk)

            return redirect('barangay-staff-document')
    else:
        form = UploadClearanceForm()

    if request.method == 'POST':
        form = UploadIndigencyForm(request.POST, request.FILES)
        if form.is_valid():
            uploaded_file = request.FILES['indigency_file']
            
            # Save the file to the user's barangay folder in the media directory
            folder_path = os.path.join(settings.MEDIA_ROOT, 'doc_templates', str(barangay.id), 'indigency')
            os.makedirs(folder_path, exist_ok=True)
            
            with open(os.path.join(folder_path, uploaded_file.name), 'wb') as destination:
                for chunk in uploaded_file.chunks():
                    destination.write(chunk)

            return redirect('barangay-staff-document')
    else:
        form = UploadIndigencyForm()

    transactions = Transaction.objects.filter(barangay=barangay).order_by('-request_date')[:10]

    context = {
        'residents': page, 
        'resident_filter': resident_filter, 
        'start_index': start_index,
        'form': form,
        'clearance_files': clearance_files, 
        'indigency_files': indigency_files, 
        'transactions': transactions,
    }
    return render(request, 'barangayStaff/document.html', context)

def edit_clearance(request, clearance_file):
    document_path = os.path.join(settings.MEDIA_ROOT, 'doc_templates', str(request.user.barangay.id), 'clearance', clearance_file)

    if not os.path.exists(document_path):
        messages.error(request, f"Document {clearance_file} not found.")
        return redirect('barangay-staff-document')

    doc = Document(document_path)

    temp_path = os.path.join(settings.MEDIA_ROOT, 'temp', 'temp_document.docx')
    doc.save(temp_path)

    with open(temp_path, 'rb') as document_file:
        response = HttpResponse(document_file.read(), content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        response['Content-Disposition'] = f'attachment; filename={clearance_file}'
        return response

def delete_clearance(request, clearance_file):
    # Add your logic for deleting a document here
    # Retrieve the document based on the doc_file parameter
    # For example:
    document_path = os.path.join(settings.MEDIA_ROOT, 'doc_templates', str(request.user.barangay.id), 'clearance', clearance_file)
    # Implement your logic for deleting the document
    if os.path.exists(document_path):
        os.remove(document_path)

    return redirect('barangay-staff-document')
    
def edit_indigency(request, indigency_file):
    document_path = os.path.join(settings.MEDIA_ROOT, 'doc_templates', str(request.user.barangay.id), 'indigency', indigency_file)

    if not os.path.exists(document_path):
        messages.error(request, f"Document {clearance_file} not found.")
        return redirect('barangay-staff-document')

    doc = Document(document_path)

    temp_path = os.path.join(settings.MEDIA_ROOT, 'temp', 'temp_document.docx')
    doc.save(temp_path)

    with open(temp_path, 'rb') as document_file:
        response = HttpResponse(document_file.read(), content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        response['Content-Disposition'] = f'attachment; filename={indigency_file}'
        return response

def delete_indigency(request, indigency_file):
    # Add your logic for deleting a document here
    # Retrieve the document based on the doc_file parameter
    # For example:
    document_path = os.path.join(settings.MEDIA_ROOT, 'doc_templates', str(request.user.barangay.id), 'indigency', indigency_file)
    # Implement your logic for deleting the document
    if os.path.exists(document_path):
        os.remove(document_path)

    return redirect('barangay-staff-document')
    
@login_required
@permission_required('barangayApp.can_view_barangay_staff_views', raise_exception=True)
def logout(request):
    return HttpResponse('Logout')


@login_required
@permission_required('barangayApp.can_view_barangay_staff_views', raise_exception=True)
def add_resident(request):
    if request.method == "POST":
        if ( 
            request.POST.get("first_name")
            and request.POST.get("last_name")
            and request.POST.get("birth_date")
            and request.POST.get("phone_number")
            and request.POST.get("sex")
            and request.POST.get("status")
            and request.POST.get("vaccination")
        ):
            resident = Resident()
            resident.first_name = request.POST.get("first_name")
            resident.last_name = request.POST.get("last_name")
            resident.middle_name = request.POST.get("middle_name")
            resident.suffix = request.POST.get("suffix")
            resident.birth_date = request.POST.get("birth_date")
            resident.phone_number = request.POST.get("phone_number")
            # resident.relationship = request.POST.get("relationship")
            resident.sex = request.POST.get("sex")
            resident.status = request.POST.get("status")
            resident.vaccination = request.POST.get("vaccination")
            
            birth_date_str = request.POST.get("birth_date")
            birth_date = datetime.strptime(birth_date_str, "%Y-%m-%d").date()
            
            resident.birth_date = birth_date
            today = timezone.now().date()
            age = today.year - birth_date.year - ((today.month, today.day) < (birth_date.month, birth_date.day))
            resident.age = age

            if request.user.user_type == 'barangay_staff':
                resident.municipal = request.user.barangay.municipal
                resident.barangay = request.user.barangay

            resident.save()
     
            health = Health()
            health.date_recorded = timezone.now().date()

            health.disability = request.POST.get("disability")
            health.illness = request.POST.get("illness")
            health.pregnancy_status = request.POST.get("pregnancy_status")
            health.malnutrition_status = request.POST.get("malnutrition_status")
            health.resident = resident
            health.save()

            economic = Economic()
            economic.specific_occupation = request.POST.get("specific_occupation")
            economic.industry_of_work = request.POST.get("industry_of_work")
            economic.working_arrangement = request.POST.get("working_arrangement")
            economic.nature_of_employment = request.POST.get("nature_of_employment")
            economic.class_of_worker = request.POST.get("class_of_worker")
            economic.basis_of_payment = request.POST.get("basis_of_payment")
            economic.resident = resident
            economic.save()

            education = Education()
            education.early_childhood_education = request.POST.get("early_childhood_education")
            education.primary_education = request.POST.get("primary_education")
            education.lower_secondary_education = request.POST.get("lower_secondary_education")
            education.upper_secondary_education = request.POST.get("upper_secondary_education")
            education.bachelor_education = request.POST.get("bachelor_education")
            education.master_doctoral_education = request.POST.get("master_doctoral_education")
            education.curriculum_strands = request.POST.get("curriculum_strands")
            education.undergraduate_reason = request.POST.get("undergraduate_reason")
            education.resident = resident
            education.save()

            
            messages.success(request, 'Resident added successfully!')
            return HttpResponseRedirect("/barangay-staff/resident")
        else:
            return render(request, "barangayStaff/add_resident.html")
    else:
        return render(request, "barangayStaff/add_resident.html")


@login_required
@permission_required('barangayApp.can_view_barangay_staff_views', raise_exception=True)
def edit_resident(request, pk):
    resident = get_object_or_404(Resident, id=pk)
    health = resident.health_set.first()
    economic = resident.economic_set.first()
    education = resident.education_set.first()

    if request.method == "POST":
        if ( 
            request.POST.get("first_name")
            and request.POST.get("last_name")
            # and request.POST.get("middle_name")
            and request.POST.get("birth_date")
            and request.POST.get("phone_number")
            # and request.POST.get("relationship")
            and request.POST.get("sex")
            and request.POST.get("status")
            and request.POST.get("vaccination")
        ):
            resident.first_name = request.POST.get("first_name")
            resident.last_name = request.POST.get("last_name")
            resident.middle_name = request.POST.get("middle_name")
            resident.suffix = request.POST.get("suffix")
            resident.birth_date = request.POST.get("birth_date")
            resident.phone_number = request.POST.get("phone_number")
            # resident.relationship = request.POST.get("relationship")
            resident.sex = request.POST.get("sex")
            resident.status = request.POST.get("status")
            resident.vaccination = request.POST.get("vaccination")
            
            birth_date_str = request.POST.get("birth_date")
            birth_date = datetime.strptime(birth_date_str, "%Y-%m-%d").date()
            
            resident.birth_date = birth_date
            today = timezone.now().date()
            age = today.year - birth_date.year - ((today.month, today.day) < (birth_date.month, birth_date.day))
            resident.age = age

            resident.save()

            health.date_recorded = timezone.now().date()   

            health.disability = request.POST.get("disability")
            health.illness = request.POST.get("illness")
            health.pregnancy_status = request.POST.get("pregnancy_status")
            health.malnutrition_status = request.POST.get("malnutrition_status")
            health.resident = resident
            health.save()

            economic.specific_occupation = request.POST.get("specific_occupation")
            economic.industry_of_work = request.POST.get("industry_of_work")
            economic.working_arrangement = request.POST.get("working_arrangement")
            economic.nature_of_employment = request.POST.get("nature_of_employment")
            economic.class_of_worker = request.POST.get("class_of_worker")
            economic.basis_of_payment = request.POST.get("basis_of_payment")
            economic.resident = resident
            economic.save()

            education.early_childhood_education = request.POST.get("early_childhood_education")
            education.primary_education = request.POST.get("primary_education")
            education.lower_secondary_education = request.POST.get("lower_secondary_education")
            education.upper_secondary_education = request.POST.get("upper_secondary_education")
            education.bachelor_education = request.POST.get("bachelor_education")
            education.master_doctoral_education = request.POST.get("master_doctoral_education")
            education.curriculum_strands = request.POST.get("curriculum_strands")
            education.undergraduate_reason = request.POST.get("undergraduate_reason")
            education.resident = resident
            education.save()

            messages.success(request, "Resident information updated successfully!")

            return HttpResponseRedirect("/barangay-staff/resident")
        else:

            messages.error(request, "Failed to update resident information. Please check the form.")

            context = {
                "resident": resident,
                "health": health,
                "economic": economic,
                "education": education,
            }
            return render(request, "barangayStaff/edit_resident.html", context)


def delete_resident(request, pk):
    resident = Resident.objects.get(id=pk)
    resident.delete()
    return redirect('barangay-staff-resident')

def tourist_spot_detail(request, id):
    spot = get_object_or_404(TouristSpot, id=id)
    return render(request, 'tourist_spots/tourist_spot_detail.html', {'spot': spot})


@login_required
@permission_required('barangayApp.can_view_barangay_staff_views', raise_exception=True)
def add_household(request):
    if request.method == 'POST':
        num_senior_citizens = request.POST.get('num_senior_citizens')
        num_pregnant_lactating_mothers = request.POST.get('num_pregnant_lactating_mothers')
        num_beneficiaries_with_disability = request.POST.get('num_beneficiaries_with_disability')
        num_registered_voters = request.POST.get('num_registered_voters')
        housing_condition = request.POST.get('housing_condition')
        monthly_income_range = request.POST.get('monthly_income_range')
        access_to_basic_amenities = request.POST.get('access_to_basic_amenities')
        ownership_of_assets = request.POST.get('ownership_of_assets')
        education_level_of_head = request.POST.get('education_level_of_head')
        household_type = request.POST.get('household_type')
        latitude = request.POST.get('latitude')
        longitude = request.POST.get('longitude')

        household = Household(
            num_senior_citizens=num_senior_citizens,
            num_pregnant_lactating_mothers=num_pregnant_lactating_mothers,
            num_beneficiaries_with_disability=num_beneficiaries_with_disability,
            num_registered_voters=num_registered_voters,
            housing_condition=housing_condition,
            monthly_income_range=monthly_income_range,
            access_to_basic_amenities=access_to_basic_amenities,
            ownership_of_assets=ownership_of_assets,
            education_level_of_head=education_level_of_head,
            household_type=household_type,
            latitude=latitude,
            longitude=longitude,
        )

        household.save()

        if request.user.user_type == 'barangay_staff':
            household.municipal = request.user.barangay.municipal
            household.barangay = request.user.barangay
        
        household.save()

        selected_head_ids = request.POST.get('selected_head_ids_adds') 
        selected_head = Resident.objects.get(id=selected_head_ids)
        household.household_head = selected_head
        
        household.save()

        selected_members_ids_adds = request.POST.get('selected_member_ids_adds')
        members_ids = selected_members_ids_adds.split(',')
        num_members = 0
        for member_id in members_ids:
            num_members += 1

        household.num_members = num_members
        household.save()
        household.members.set(members_ids)


        messages.success(request, 'Household added successfully!')
        return redirect('barangay-staff-household') 

    return render(request, 'barangayStaff/add_household.html')


@login_required
@permission_required('barangayApp.can_view_barangay_staff_views', raise_exception=True)
def edit_household(request, pk):
    household = get_object_or_404(Household, id=pk)
    error_message = None
    if request.method == "POST":
        if ( 
            request.POST.get("selected_head_ids_edit")
            and request.POST.get("selected_member_ids_edit")
            and request.POST.get("num_senior_citizens")
            and request.POST.get("num_pregnant_lactating_mothers")
            and request.POST.get("num_beneficiaries_with_disability")
            and request.POST.get("num_registered_voters")
            and request.POST.get("monthly_income_range")
            and request.POST.get("access_to_basic_amenities")
            and request.POST.get("ownership_of_assets")
            and request.POST.get("education_level_of_head")
            and request.POST.get("housing_condition")
            and request.POST.get("household_type")
            and request.POST.get("latitude")
            and request.POST.get("longitude")
        ):
            selected_head_id = request.POST.get("selected_head_ids_edit")
            selected_head = Resident.objects.get(id=selected_head_id)
            household.household_head = selected_head

            selected_member_ids = request.POST.get('selected_member_ids_edit')
            members_ids = selected_member_ids.split(',')
            num_members = 0
            for member_id in members_ids:
                num_members += 1
            household.num_members = num_members
            household.members.set(members_ids)

            household.num_senior_citizens = request.POST.get("num_senior_citizens")
            household.num_pregnant_lactating_mothers = request.POST.get("num_pregnant_lactating_mothers")
            household.num_beneficiaries_with_disability = request.POST.get("num_beneficiaries_with_disability")
            household.num_registered_voters = request.POST.get("num_registered_voters")
            household.monthly_income_range = request.POST.get("monthly_income_range")
            household.access_to_basic_amenities = request.POST.get("access_to_basic_amenities")
            household.ownership_of_assets = request.POST.get("ownership_of_assets")
            household.education_level_of_head = request.POST.get("education_level_of_head")
            household.housing_condition = request.POST.get("housing_condition")
            household.household_type = request.POST.get("household_type")
            household.latitude = request.POST.get("latitude")
            household.longitude = request.POST.get("longitude")


            household.save()

            return HttpResponseRedirect("/barangay-staff/household")
        else:
            error_message = "Please fill out all required fields."

        context = {
            'household': household,
            'error_message': error_message
        }
    return render(request, "barangayStaff/edit_household.html", context)

def search_members(request):
    user = request.user
    barangay = user.barangay

    if not barangay:
        return HttpResponse("You are not associated with any barangay.")

    query = request.GET.get('q')

    matching_residents = Resident.objects.filter(
        models.Q(first_name__icontains=query) |
        models.Q(last_name__icontains=query) |
        models.Q(middle_name__icontains=query) |
        models.Q(suffix__icontains=query),
        barangay=barangay  # Filter residents based on the user's barangay
    )

    matching_members = [
        {
            'id': resident.id,
            'full_name': f'{resident.last_name}, {resident.first_name} {resident.middle_name}  {resident.suffix}',
        }
        for resident in matching_residents
    ]

    results = [member for member in matching_members if query.lower() in member['full_name'].lower()]
    
    return JsonResponse({'results': results})


def search_residents(request):
    user = request.user
    barangay = user.barangay

    if not barangay:
        return HttpResponse("You are not associated with any barangay.")

    query = request.GET.get('q')

    matching_residents = Resident.objects.filter(
        models.Q(first_name__icontains=query) |
        models.Q(last_name__icontains=query) |
        models.Q(middle_name__icontains=query) |
        models.Q(suffix__icontains=query),
        barangay=barangay  # Filter residents based on the user's barangay
    )

    matching_residents = [
        {
            'id': resident.id,
            'full_name': f'{resident.first_name} {resident.middle_name} {resident.last_name} {resident.suffix}',
        }
        for resident in matching_residents
    ]

    results = list(matching_residents)
    return JsonResponse({'results': results})

def get_household_data(request):
    if request.method == 'GET':
        household_id = request.GET.get('household_id')

        try:
            household = Household.objects.get(id=household_id)
            household_head = household.household_head
            household_members = household.members.all()

            head_data = {
                'full_name': household_head.full_name() if household_head else None,
                'id': household_head.id if household_head else None
            }

            members_data = [
                {
                    'full_name': member.full_name(),
                    'id': member.id
                }
                for member in household_members
            ]

            household_data = {
                'head': head_data,
                'members': members_data
            }
            return JsonResponse(household_data)
        except Household.DoesNotExist:
            return JsonResponse({'error': 'Household not found'}, status=404)

    return JsonResponse({'error': 'Invalid request'}, status=400)
    
class ExportHouseholdsPDFView(View):
    def get(self, request):
        user = request.user
        barangay = user.barangay
        if not barangay:
            return HttpResponse("You are not associated with any barangay.")

        # Retrieve filter parameters from the URL query string
        filter_params = {k: v for k, v in request.GET.items() if v}  # Remove empty values

        households = Household.objects.filter(barangay=barangay).order_by('household_head__last_name')
        household_filter = HouseholdFilters(filter_params, queryset=households)
        filtered_households = household_filter.qs

        response = HttpResponse(content_type='application/pdf')
        response['Content-Disposition'] = 'attachment; filename="household_report.pdf"'

        pdf = SimpleDocTemplate(response, pagesize=letter)
        elements = []

        # Define the default and filterable fields
        default_fields = ['#', 'Household Head', 'Members']
        filterable_fields = [
            ('num_senior_citizens', 'Number of Senior Citizens'),
            ('num_pregnant_lactating_mothers', 'Numbers of Pregnant\Lactating Mothers'),
            ('num_beneficiaries_with_disability', 'Numbers of Members with Disability'),
            ('num_registered_voters', 'Registered Voters'),
            ('housing_condition', 'Household Condition'),
            ('monthly_income_range', 'Monthly Income Range'),
            ('access_to_basic_amenities', 'Access To Basic Amenities'),
            ('ownership_of_assets', 'Ownership of Assets'),
            ('education_level_of_head', 'Head Education Level'),
            ('household_type', 'Household Type'),
        ]

        # Create table headers based on filter parameters
        table_headers = [field[1] for field in filterable_fields if field[0] in filter_params]
        table_headers = default_fields + table_headers

        table_data = [table_headers]  # Header row

        for index, household in enumerate(filtered_households, start=1):
            row = [
                index,
                household.household_head.full_name() if household.household_head else None,
                # Create a paragraph for each member's name
                Paragraph('<br/>'.join([member.full_name() for member in household.members.all()]), getSampleStyleSheet()['Normal']),
            ]

            for field, label in filterable_fields:
                if field in filter_params:
                    row.append(getattr(household, field))

            table_data.append(row)

        table = Table(table_data)
        table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.green),  # Header row background color
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),  # Header row text color
            ('ALIGN', (0, 0), (-1, -1), 'CENTER'),  # Centered alignment for all cells
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),  # Header row font style
            ('FONTSIZE', (0, 0), (-1, 0), 12),  # Header row font size
            ('BOTTOMPADDING', (0, 0), (-1, 0), 12),  # Header row bottom padding
            ('BACKGROUND', (0, 1), (-1, -1), colors.white),  # Data rows background color
            ('GRID', (0, 0), (-1, -1), 1, colors.black),
        ]))

        elements.append(table)

        pdf.build(elements)

        return response

class ExportHouseholdsExcelView(View):
    def get(self, request):
        user = request.user
        barangay = user.barangay
        if not barangay:
            return HttpResponse("You are not associated with any barangay.")

        # Retrieve filter parameters from the URL query string
        filter_params = {k: v for k, v in request.GET.items() if v}

        households = Household.objects.filter(barangay=barangay).order_by('household_head__last_name')
        household_filter = HouseholdFilters(filter_params, queryset=households)
        filtered_households = household_filter.qs

        response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')
        response['Content-Disposition'] = 'attachment; filename="household_report.xlsx"'

        workbook = Workbook()
        worksheet = workbook.active

        default_fields = ['Household Head', 'Household Members']
        filterable_fields = [
            ('num_senior_citizens', 'Number of Senior Citizens'),
            ('num_pregnant_lactating_mothers', 'Numbers of Pregnant\Lactating Mothers'),
            ('num_beneficiaries_with_disability', 'Numbers of Members with Disability'),
            ('num_registered_voters', 'Registered Voters'),
            ('housing_condition', 'Household Condition'),
            ('monthly_income_range', 'Monthly Income Range'),
            ('access_to_basic_amenities', 'Access To Basic Amenities'),
            ('ownership_of_assets', 'Ownership of Assets'),
            ('education_level_of_head', 'Head Education Level'),
            ('household_type', 'Household Type'),
            ('housing_condition', 'Household Condition'),
        ]

        # Create table headers based on filter parameters
        table_headers = [field[1] for field in filterable_fields if field[0] in filter_params]
        table_headers = default_fields + table_headers

        for col_num, header in enumerate(table_headers, 1):
            col_letter = get_column_letter(col_num)
            worksheet[f'{col_letter}1'] = header
            worksheet.column_dimensions[col_letter].width = len(header) + 2

        for index, household in enumerate(filtered_households, start=2):
            members_names = ', '.join([member.full_name() for member in household.members.all()])

            row_data = [
                household.household_head.full_name() if household.household_head else None,
                members_names,
            ]

            for field, label in filterable_fields:
                if field in filter_params:
                    row_data.append(getattr(household, field))

            for col_num, cell_value in enumerate(row_data, 1):
                col_letter = get_column_letter(col_num)
                worksheet[f'{col_letter}{index}'] = cell_value

        workbook.save(response)

        return response

class GeneratePDF(View):
    def get(self, request):
        user = request.user
        barangay = user.barangay
        if not barangay:
            return HttpResponse("You are not associated with any barangay.")
            
        residents = Resident.objects.filter(barangay=barangay).order_by('last_name')
        # residents = Resident.objects.all()
        resident_filter = ResidentFilters(request.GET, queryset=residents)
        filtered_residents = resident_filter.qs

        response = HttpResponse(content_type='application/pdf')
        response['Content-Disposition'] = 'attachment; filename="resident_report.pdf"'

        pdf = SimpleDocTemplate(response, pagesize=letter)
        elements = []

        table_data = [['#', 'Name', 'Age', 'Status', 'Birthdate']]  # Header row

        # for index, household in enumerate(filtered_households, start=1):
        for index, resident in enumerate(filtered_residents, start=1):
            row = [index, f"{resident.last_name}, {resident.first_name} {resident.middle_name} {resident.suffix}", 
                    resident.age, 
                    resident.status, 
                    resident.birth_date]

            table_data.append(row)

        table = Table(table_data)
        table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.green),  # Header row background color
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),  # Header row text color
            ('ALIGN', (0, 0), (-1, -1), 'CENTER'),  # Centered alignment for all cells
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),  # Header row font style
            ('FONTSIZE', (0, 0), (-1, 0), 12),  # Header row font size
            ('BOTTOMPADDING', (0, 0), (-1, 0), 12),  # Header row bottom padding
            ('BACKGROUND', (0, 1), (-1, -1), colors.white),  # Data rows background color
            ('GRID', (0, 0), (-1, -1), 1, colors.black),
            ('ALIGN', (1, 1), (-1, -1), 'LEFT'),  # Left-align the cell containing the name
        ]))


        elements.append(table)

        pdf.build(elements)

        return response

class GenerateExcel(View):
    def get(self, request):
        user = request.user
        barangay = user.barangay
        if not barangay:
            return HttpResponse("You are not associated with any barangay.")
            
        residents = Resident.objects.filter(barangay=barangay).order_by('last_name')
        resident_filter = ResidentFilters(request.GET, queryset=residents)
        filtered_residents = resident_filter.qs

        response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')
        response['Content-Disposition'] = 'attachment; filename="resident_report.xlsx"'

        workbook = Workbook()
        worksheet = workbook.active

        # Define the filterable fields
        filterable_fields = ['Name', 'Age', 'Vaccination', 'Status', 'Birthdate']

        # Create table headers based on filter parameters
        table_headers = filterable_fields

        for col_num, header in enumerate(table_headers, 1):
            col_letter = get_column_letter(col_num)
            worksheet[f'{col_letter}1'] = header
            worksheet.column_dimensions[col_letter].width = len(header) + 2

        for index, resident in enumerate(filtered_residents, start=2):
            row_data = [
                f"{resident.last_name}, {resident.first_name} {resident.middle_name} {resident.suffix}",
                resident.age,
                resident.vaccination,
                resident.status,
                resident.birth_date,
            ]

            for col_num, cell_value in enumerate(row_data, 1):
                col_letter = get_column_letter(col_num)
                worksheet[f'{col_letter}{index}'] = cell_value

        workbook.save(response)

        return response

def generate_clearance(request, resident_id):
    # Get the resident object
    resident = Resident.objects.get(id=resident_id)

    # Get the path to the folder containing docx files for clearance based on the barangay
    clearance_folder_path = os.path.join(settings.MEDIA_ROOT, 'doc_templates', str(resident.barangay.id), 'clearance')

    # Get the list of docx files in the folder
    docx_files = [f for f in os.listdir(clearance_folder_path) if f.endswith('.docx')]

    # Choose the appropriate docx file (you might need to implement a logic to select the right one)
    # For now, let's assume we choose the first file in the list
    if docx_files:
        clearance_template_path = os.path.join(clearance_folder_path, docx_files[0])
    else:
        return HttpResponse("No clearance template found for this barangay.")

    # Load the Word template
    doc = Document(clearance_template_path)

    current_date = date.today()

    day = current_date.strftime('%d')
    month = current_date.strftime('%B')
    year = current_date.strftime('%Y')

    data_mapping = {
        '[resident_name]': f'{resident.first_name} {resident.middle_name} {resident.last_name} {resident.suffix}',
        '[age]': str(resident.age),
        '[sex]': resident.sex,
        '[birth_date]': str(resident.birth_date),
        '[phone_number]': resident.phone_number,
        '[vaccination]': resident.vaccination,
        '[status]': resident.status,
        '[barangay_name]': resident.barangay.name,
        '[municipality_name]': resident.municipal.name,
        '[issue_date_day]': day,
        '[issue_date_month]': month,
        '[issue_date_year]': year,
        '[official_signature]': 'John Doe',
        '[official_name]': 'John Doe',
        '[official_position]': 'Barangay Official',
        '[official_barangay]': 'Barangay ABC',
    }

    # Iterate through paragraphs in the document
    for para in doc.paragraphs:
        for key, value in data_mapping.items():
            if key in para.text:
                para.text = para.text.replace(key, value)

    # Save the modified document to the response
    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = f'attachment; filename=barangay_clearance_{resident.last_name}_{resident.first_name}.docx'
    doc.save(response)

    # Log the transaction
    transaction = Transaction.objects.create(
        barangay=resident.barangay,
        requester=resident,  # Use the resident who requested the clearance
        request_date=datetime.now().strftime('%Y-%m-%d %H:%M:%S'), 
        document_type='barangay_clearance',
        # Add other necessary fields for the Transaction model
    )

    return response

def generate_indigency(request, resident_id):
    # Get the resident object
    resident = Resident.objects.get(id=resident_id)

    # Get the path to the folder containing docx files for indigency based on the barangay
    indigency_folder_path = os.path.join(settings.MEDIA_ROOT, 'doc_templates', str(resident.barangay.id), 'indigency')

    # Get the list of docx files in the folder
    docx_files = [f for f in os.listdir(indigency_folder_path) if f.endswith('.docx')]

    # Choose the appropriate docx file (you might need to implement a logic to select the right one)
    # For now, let's assume we choose the first file in the list
    if docx_files:
        indigency_template_path = os.path.join(indigency_folder_path, docx_files[0])
    else:
        return HttpResponse("No indigency template found for this barangay.")

    # Load the Word template
    doc = Document(indigency_template_path)

    current_date = date.today()

    day = current_date.strftime('%d')
    month = current_date.strftime('%B')
    year = current_date.strftime('%Y')

    data_mapping = {
        '[resident_name]': f'{resident.first_name} {resident.middle_name} {resident.last_name} {resident.suffix}',
        '[issue_date_day]': day,
        '[issue_date_month]': month,
        '[issue_date_year]': year,
    }

    # Iterate through paragraphs in the document
    for para in doc.paragraphs:
        for key, value in data_mapping.items():
            if key in para.text:
                para.text = para.text.replace(key, value)

    # Save the modified document to the response
    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = f'attachment; filename=barangay_indigency_{resident.last_name}_{resident.first_name}.docx'
    doc.save(response)

    # Log the transaction
    transaction = Transaction.objects.create(
        barangay=resident.barangay,
        requester=resident,
        request_date=datetime.now().strftime('%Y-%m-%d %H:%M:%S'), 
        document_type='certificate_of_indigency',
    )

    return response

def transactions(request):
    user = request.user
    barangay = user.barangay
    if not barangay:
        return HttpResponse("You are not associated with any barangay.")

    transactions = Transaction.objects.filter(barangay=barangay).order_by('-request_date')

    p = Paginator(transactions, 10)  
    page_num = request.GET.get('page', 1)
    page = p.page(page_num)
    start_index = (page.number - 1) * p.per_page

    context = {
        'transactions': page,
        'start_index': start_index
    }

    return render(request, 'barangayStaff/transactions.html', context)